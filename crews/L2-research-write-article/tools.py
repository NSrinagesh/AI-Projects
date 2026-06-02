import os
from crewai.tools import tool
from docx import Document
from docx.shared import Pt

@tool("Save Article as DOCX")
def save_as_docx(markdown_text: str) -> str:
    """
    Converts a markdown-formatted article into a formatted Word document (.docx)
    and saves it to the output folder. Use this tool when the final edited
    article is ready for publication. Pass the full markdown text as input.
    """
    os.makedirs("output", exist_ok=True)
    doc = Document()

    for line in markdown_text.split("\n"):
        line = line.strip()

        if not line:
            continue
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            paragraph = doc.add_paragraph(line)
            paragraph.style.font.size = Pt(11)

    output_path = "output/article.docx"
    doc.save(output_path)
    return f"Article saved as Word document at {output_path}"
